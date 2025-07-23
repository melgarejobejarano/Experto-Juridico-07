from __future__ import annotations

import json
import os
import tempfile
import time
import csv  # Agregamos la importación de csv
from datetime import datetime
from pathlib import Path
from typing import Literal

import openai
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader

###############################################################################
# Configuración de la página y tema                                           
###############################################################################

import base64
import requests
from io import BytesIO
from PIL import Image

# Colores corporativos
COLORS = {
    "primary": "#2D5DA0",    # Azul corporativo
    "secondary": "#E5B448",  # Dorado complementario
    "neutral": "#718096",    # Gris neutro
    "background": "#F8FAFC", # Fondo claro
    "text": "#1A202C",       # Texto oscuro
    "light_primary": "#EEF2F7", # Azul muy claro para franjas
    "footer": "#1E3A5F",     # Azul oscuro para pie de página
    "button_text": "#F8FAFC", # Texto claro para botones
    "button_text_hover": "#FFFFFF", # Texto blanco puro para hover
    "star_inactive": "#D1D5DB", # Gris para estrellas inactivas
    "star_active": "#F59E0B", # Dorado para estrellas activas
    "star_hover": "#FBBF24",  # Dorado claro para hover
}

def generar_imagen_corporativa():
    """Genera una imagen corporativa usando DALL-E."""
    image_path = os.path.join(os.path.dirname(__file__), "corporate_image.png")
    
    # Si la imagen ya existe, la retornamos
    if os.path.exists(image_path):
        return image_path
        
    try:
        response = openai.images.generate(
            model="dall-e-3",
            prompt="Ejecutivo de negocios con traje oscuro interactuando con una interfaz holográfica azul; en el centro, un icono luminoso de balanza de la justicia rodeado de símbolos legales —mazo, edificio judicial, apretón de manos, lupa, grupo de personas, libro abierto y globo de diálogo— sobre un fondo oscuro difuminado que realza el efecto futurista y profesional.",
            n=1,
            size="1024x1024",
            quality="hd",
            style="vivid"
        )
        
        image_url = response.data[0].url
        image_data = requests.get(image_url).content
        
        with open(image_path, "wb") as f:
            f.write(image_data)
            
        return image_path
    except Exception as e:
        st.error(f"Error generando la imagen corporativa: {str(e)}")
        fallback_image = os.path.join(os.path.dirname(__file__), "abogada experta japon.png")
        return fallback_image if os.path.exists(fallback_image) else None

# Configuración de la página
st.set_page_config(
    page_title="Experto Jurídico 7.0",
    page_icon="⚖️",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Inyección de CSS personalizado
st.markdown(
    f"""
    <style>
        /* Tipografía y colores base */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');
        
        :root {{
            --primary-color: {COLORS["primary"]};
            --secondary-color: {COLORS["secondary"]};
            --neutral-color: {COLORS["neutral"]};
            --background-color: {COLORS["background"]};
            --text-color: {COLORS["text"]};
            --light-primary: {COLORS["light_primary"]};
        }}
        
        .stApp {{
            font-family: 'Inter', sans-serif;
            background-color: var(--background-color);
        }}
        
        /* Ocultar elementos no deseados */
        #MainMenu {{display: none;}}
        footer {{display: none;}}
        
        /* Contenedor principal más compacto */
        .main .block-container {{
            max-width: 75%;  /* Reducido al 75% del ancho */
            padding: 2rem 3rem;
            margin: 0 auto;
            background-color: white;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }}
        
        /* Franja descriptiva */
        .descriptive-band {{
            background-color: var(--light-primary);
            padding: 1rem;
            border-radius: 8px;
            margin: 1rem 0;
            text-align: center;
            color: var(--primary-color);
        }}
        
        /* Info box */
        .info-box {{
            background-color: rgba(45, 93, 160, 0.05);
            border-left: 3px solid var(--primary-color);
            padding: 1rem;
            margin: 1rem 0;
            border-radius: 0 8px 8px 0;
        }}
        
        /* Contacto */
        .contact-info {{
            background-color: white;
            padding: 1.5rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
            margin: 1rem 0;
        }}
        
        .contact-info h4 {{
            color: var(--primary-color);
            margin-bottom: 1rem;
        }}
        
        .contact-detail {{
            display: flex;
            align-items: center;
            gap: 0.5rem;
            margin: 0.5rem 0;
            color: var(--text-color);
        }}
        
        /* Navegación superior */
        .stNavigationContainer {{
            background-color: white;
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            z-index: 100;
            padding: 1rem;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        
        /* Botones */
        .stButton > button {{
            width: 100%;
            border-radius: 8px;
            font-weight: 500;
            transition: all 0.2s;
        }}
        
        .stButton > button:first-child {{
            background-color: var(--primary-color);
            color: white;
        }}
        
        .stButton > button:not(:first-child) {{
            background-color: var(--neutral-color);
            color: white;
        }}
        
        /* Radio buttons en línea */
        .stRadio > div {{
            display: flex;
            gap: 1rem;
            flex-wrap: wrap;
        }}
        
        .stRadio label {{
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }}
        
        /* Paneles desplegables */
        .streamlit-expanderHeader {{
            background-color: white;
            border-radius: 8px;
            padding: 1rem;
            margin-bottom: 1rem;
        }}
        
        /* Modo oscuro */
        @media (prefers-color-scheme: dark) {{
            :root {{
                --background-color: #1A202C;
                --text-color: #F8FAFC;
            }}
            
            .stApp {{
                background-color: var(--background-color);
                color: var(--text-color);
            }}
        }}
        
        /* Accesibilidad */
        *:focus {{
            outline: 3px solid var(--primary-color);
            outline-offset: 2px;
        }}
        
        /* Spinner personalizado */
        .stSpinner {{
            border-color: var(--primary-color);
        }}
        
        /* Navegación superior mejorada */
        .nav-container {{
            display: grid;
            grid-template-columns: repeat(6, 1fr);
            gap: 0.5rem;
            padding: 1rem;
            background-color: white;
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            z-index: 100;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            max-width: 1200px;
            margin: 0 auto;
        }}
        
        /* Botones de navegación */
        .stButton {{
            width: 100%;
        }}
        
        .stButton > button {{
            width: 100% !important;
            min-width: 0 !important;
            padding: 0.75rem 0.5rem !important;
            background-color: var(--primary-color) !important;
            border: none !important;
            color: {COLORS["button_text"]} !important;
            font-size: 0.875rem !important;
            font-weight: 500 !important;
            line-height: 1.2 !important;
            border-radius: 8px !important;
            transition: all 0.2s ease !important;
            height: auto !important;
            white-space: nowrap !important;
            overflow: hidden !important;
            text-overflow: ellipsis !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            gap: 0.25rem !important;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
        }}
        
        .stButton > button:hover {{
            background-color: {COLORS["primary"]}DD !important;
            color: {COLORS["button_text_hover"]} !important;
            box-shadow: 0 4px 8px rgba(0,0,0,0.15) !important;
            transform: translateY(-1px) !important;
        }}
        
        .stButton > button:active {{
            transform: translateY(1px) !important;
            box-shadow: inset 0 2px 4px rgba(0,0,0,0.1) !important;
        }}
        
        /* Contenedor principal ajustado para la navegación */
        .main .block-container {{
            padding-top: 5rem !important;  /* Espacio para la navegación fija */
        }}
        
        /* Media queries para responsividad */
        @media (max-width: 768px) {{
            .nav-container {{
                grid-template-columns: repeat(3, 1fr);
                padding: 0.75rem;
            }}
            
            .stButton > button {{
                padding: 0.5rem 0.25rem !important;
                font-size: 0.75rem !important;
            }}
        }}
        
        /* Sistema de estrellas */
        .rating {{
            display: inline-flex;
            gap: 0.25rem;
            padding: 0.5rem;
            border-radius: 8px;
        }}
        
        .star {{
            font-size: 2rem;
            cursor: pointer;
            color: #ddd;
            transition: color 0.2s ease;
        }}
        
        .star.filled {{
            color: {COLORS["secondary"]};
        }}
        
        .star:hover,
        .star:focus {{
            color: {COLORS["secondary"]};
            outline: none;
        }}
        
        /* Pie de página personalizado */
        .custom-footer {{
            position: fixed;
            bottom: 0;
            left: 0;
            right: 0;
            background-color: {COLORS["footer"]};
            color: white;
            padding: 1rem;
            text-align: center;
            font-size: 0.9rem;
            z-index: 100;
        }}
        
        .custom-footer a {{
            color: white;
            text-decoration: none;
        }}
        
        /* Ajuste para el contenido principal */
        .main .block-container {{
            margin-bottom: 4rem;  /* Espacio para el footer */
        }}
        
        /* Contador de estrellas */
        .rating-count {{
            margin-left: 1rem;
            font-size: 1.1rem;
            color: var(--text-color);
            display: inline-flex;
            align-items: center;
            padding: 0.25rem 0.75rem;
            background-color: var(--light-primary);
            border-radius: 1rem;
        }}
        
        /* Navegación superior mejorada */
        .nav-container {{
            display: flex;
            gap: 0.5rem;
            padding: 1rem;
            background-color: white;
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            z-index: 100;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        
        .nav-button {{
            flex: 1;
            background-color: var(--primary-color);
            color: {COLORS["button_text"]} !important;
            padding: 0.75rem 1rem;
            border-radius: 8px;
            border: none;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            transition: all 0.2s ease;
            font-weight: 500;
            text-align: center;
            cursor: pointer;
            text-decoration: none;
            display: flex;
            align-items: center;
            justify-content: center;
            min-width: 120px;
            font-size: 0.95rem;
            letter-spacing: 0.02em;
        }}
        
        .nav-button:hover,
        .nav-button:focus,
        .nav-button:active {{
            background-color: {COLORS["primary"]}DD;
            color: {COLORS["button_text_hover"]} !important;
            box-shadow: 0 4px 8px rgba(0,0,0,0.15);
            transform: translateY(-1px);
            text-decoration: none;
        }}
        
        /* Asegurar que los emojis en los botones tengan buen contraste */
        .nav-button span {{
            filter: brightness(1.1);
            margin-right: 0.5rem;
        }}
        
        /* Estilo para botón activo */
        .nav-button.active {{
            background-color: {COLORS["primary"]};
            color: {COLORS["button_text_hover"]} !important;
            box-shadow: inset 0 2px 4px rgba(0,0,0,0.1);
            transform: translateY(1px);
        }}
        
        /* Contenedor de imagen corporativa */
        .corporate-image-container {{
            width: 100%;
            margin: 1rem 0;
            border-radius: 12px;
            overflow: hidden;
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }}
        
        .corporate-image-container img {{
            width: 100%;
            height: auto;
            object-fit: cover;
            display: block;
        }}
        
        /* Botón de historial */
        .history-button {{
            display: block;
            width: fit-content;
            margin: 2rem auto;
            padding: 1rem 2rem;
            background-color: white;
            color: var(--primary-color);
            border: 2px solid var(--primary-color);
            border-radius: 8px;
            font-weight: 500;
            transition: all 0.2s ease;
            cursor: pointer;
            text-align: center;
        }}
        
        .history-button:hover {{
            background-color: var(--primary-color);
            color: white;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        }}
        
        /* Media queries para responsividad */
        @media (max-width: 768px) {{
            .nav-container {{
                flex-wrap: wrap;
            }}
            
            .nav-button {{
                min-width: calc(33.33% - 0.5rem);
            }}
        }}

        /* Sistema de valoración por estrellas */
        .rating-container {{
            display: flex;
            align-items: center;
            gap: 2rem;
            padding: 1.5rem;
            background-color: white;
            border-radius: 12px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }}
        
        .stars-container {{
            display: flex;
            gap: 0.75rem;
        }}
        
        .star-rating {{
            display: inline-flex;
            gap: 0.75rem;
        }}
        
        .star-btn {{
            background: none;
            border: none;
            cursor: pointer;
            font-size: 2.5rem;
            padding: 0.25rem;
            transition: all 0.3s ease;
            color: {COLORS["star_inactive"]};
            position: relative;
            outline: none;
        }}
        
        .star-btn:hover ~ .star-btn {{
            color: {COLORS["star_inactive"]};
        }}
        
        .star-rating:hover .star-btn {{
            color: {COLORS["star_hover"]};
        }}
        
        .star-rating .star-btn:hover ~ .star-btn {{
            color: {COLORS["star_inactive"]};
        }}
        
        .star-btn.active {{
            color: {COLORS["star_active"]};
            transform: scale(1.1);
        }}
        
        .star-btn.active ~ .star-btn {{
            color: {COLORS["star_inactive"]};
        }}
        
        .rating-value {{
            font-size: 1.1rem;
            font-weight: 500;
            color: {COLORS["text"]};
            background-color: {COLORS["light_primary"]};
            padding: 0.75rem 1.25rem;
            border-radius: 2rem;
            min-width: 120px;
            text-align: center;
        }}
    </style>
    
    <div class="custom-footer">
        Soporte técnico: <a href="mailto:melgarejobejarano@gmail.com">melgarejobejarano@gmail.com</a> | +51 961 277 144 · © 2024 Experto Jurídico
    </div>
    """,
    unsafe_allow_html=True,
)

###############################################################################
# Componentes de la interfaz                                                  
###############################################################################

def mostrar_navegacion():
    """Muestra la barra de navegación superior."""
    cols = st.columns(6)
    nav_items = [
        ("inicio", "🏠", "Inicio"),
        ("generar", "📝", "Generar"),
        ("historial", "📋", "Historial"),
        ("ayuda", "❓", "Ayuda"),
        ("contacto", "📞", "Contacto"),
        ("feedback", "⭐", "Feedback")
    ]
    
    for i, (page, icon, label) in enumerate(nav_items):
        with cols[i]:
            if st.button(
                f"{icon} {label}",
                key=f"nav_{page}",
                use_container_width=True,
                type="primary" if page == st.session_state.page else "secondary"
            ):
                # Limpiar estados si cambiamos de página
                if page != st.session_state.page:
                    if page != "generar":
                        for key in ['paso_actual', 'area', 'rol', 'analysis', 'document_text']:
                            if key in st.session_state:
                                del st.session_state[key]
                    st.session_state.page = page
                    st.rerun()

def sistema_estrellas(key_prefix=""):
    """Componente de calificación por estrellas."""
    # Inicializar el estado si no existe
    if f'rating_{key_prefix}' not in st.session_state:
        st.session_state[f'rating_{key_prefix}'] = 0
    if 'total_ratings' not in st.session_state:
        st.session_state.total_ratings = 0
        st.session_state.sum_ratings = 0

    # Crear el contenedor para las estrellas y la puntuación media
    col1, col2 = st.columns([3, 2])
    
    with col1:
        stars_html = ""
        for i in range(1, 6):
            filled = "filled" if i <= st.session_state[f"rating_{key_prefix}"] else ""
            stars_html += (
                f'<span class="star {filled}" '
                f'role="button" tabindex="0" '
                f'onclick="document.dispatchEvent(new CustomEvent(\'star-click\', {{detail: {i}}}));" '
                f'onkeypress="if(event.key===\'Enter\') document.dispatchEvent(new CustomEvent(\'star-click\', {{detail: {i}}}));">'
                f'★</span>'
            )
        
        st.markdown(
            f"""
            <div class="rating" role="group" aria-label="Calificación por estrellas">
                {stars_html}
            </div>
            """,
            unsafe_allow_html=True
        )
    
    with col2:
        if st.session_state.total_ratings > 0:
            average_rating = st.session_state.sum_ratings / st.session_state.total_ratings
            st.markdown(
                f"""
                <div class="rating-count">
                    {average_rating:.1f} ⭐ ({st.session_state.total_ratings} votos)
                </div>
                """,
                unsafe_allow_html=True
            )

def mostrar_feedback():
    """Muestra el formulario de feedback."""
    st.write("### Valoración del servicio")
    
    # Inicializar estados si no existen
    if "rating_value" not in st.session_state:
        st.session_state["rating_value"] = 0
    if "feedback_enviado" not in st.session_state:
        st.session_state["feedback_enviado"] = False
    
    # Estilos personalizados
    st.markdown(
        """
        <style>
        .rating-container {
            background-color: white;
            padding: 1.5rem;
            border-radius: 8px;
            text-align: center;
            margin: 1rem 0;
        }
        
        .stars {
            display: inline-flex;
            gap: 0.75rem;
            margin: 1rem 0;
        }
        
        .star {
            font-size: 3rem;
            cursor: pointer;
            color: #DDD;
            transition: color 0.2s ease;
        }
        
        .star.active {
            color: #FFD700;
        }
        
        .rating-value {
            font-size: 1.2rem;
            color: #666;
            margin-top: 1rem;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # Crear columnas para centrar el contenido
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        # Sistema de estrellas
        stars_html = "".join([
            f'<span class="star{" active" if i <= st.session_state["rating_value"] else ""}" '
            f'onclick="selectRating({i})">★</span>'
            for i in range(1, 6)
        ])
        
        rating_text = (f"{st.session_state['rating_value']} de 5 estrellas" 
                      if st.session_state["rating_value"] > 0 
                      else "Sin valorar")
        
        st.markdown(
            f"""
            <div class="rating-container">
                <div class="stars">
                    {stars_html}
                </div>
                <div class="rating-value">
                    {rating_text}
                </div>
            </div>
            
            <script>
                function selectRating(value) {{
                    const stars = document.querySelectorAll('.star');
                    stars.forEach((star, index) => {{
                        if (index < value) {{
                            star.classList.add('active');
                            star.style.color = '#FFD700';
                        }} else {{
                            star.classList.remove('active');
                            star.style.color = '#DDD';
                        }}
                    }});
                    
                    window.parent.postMessage({{
                        type: 'streamlit:setComponentValue',
                        value: value
                    }}, '*');
                }}
                
                // Inicializar estrellas
                const currentRating = {st.session_state["rating_value"]};
                if (currentRating > 0) {{
                    selectRating(currentRating);
                }}
            </script>
            """,
            unsafe_allow_html=True
        )
        
        # Componente oculto para manejar el estado
        rating = st.number_input(
            "Valoración en estrellas",  # Agregamos un label descriptivo
            min_value=0,
            max_value=5,
            value=st.session_state["rating_value"],
            label_visibility="collapsed"  # Ocultamos el label pero mantenemos la accesibilidad
        )
        
        if rating != st.session_state["rating_value"]:
            st.session_state["rating_value"] = rating
            st.rerun()
        
        # Área de comentarios
        comentario = st.text_area(
            "Comentarios y sugerencias",  # Label más descriptivo
            height=80,
            placeholder="¿Qué te pareció el servicio? (opcional)"
        )
        
        # Botón de envío
        if st.button(
            "Enviar valoración",
            type="primary",
            use_container_width=True,
            disabled=st.session_state["rating_value"] == 0 or st.session_state["feedback_enviado"]
        ):
            # Guardar feedback
            feedback = {
                "fecha": datetime.now().isoformat(),
                "rating": st.session_state["rating_value"],
                "comentario": comentario
            }
            
            # Guardar en CSV
            csv_path = "feedbacks.csv"
            is_new_file = not os.path.exists(csv_path)
            
            with open(csv_path, mode='a', newline='', encoding='utf-8') as file:
                writer = csv.DictWriter(file, fieldnames=["fecha", "rating", "comentario"])
                if is_new_file:
                    writer.writeheader()
                writer.writerow(feedback)
            
            # Mostrar mensaje de éxito
            st.session_state["feedback_enviado"] = True
            st.success("¡Gracias por tu valoración! 🌟")
        
        # Botón para nueva valoración
        if st.session_state["feedback_enviado"]:
            if st.button("Nueva valoración"):
                st.session_state["rating_value"] = 0
                st.session_state["feedback_enviado"] = False
                st.rerun()
    
    # Mostrar estadísticas
    if os.path.exists("feedbacks.csv"):
        with open("feedbacks.csv", mode='r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            feedbacks = list(reader)
        
        if feedbacks:
            st.markdown("---")
            total_ratings = len(feedbacks)
            avg_rating = sum(float(f["rating"]) for f in feedbacks) / total_ratings
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Promedio", f"{avg_rating:.1f} ⭐")
            with col2:
                st.metric("Total valoraciones", str(total_ratings))

# Asegurarnos de que los estilos CSS no hagan referencia a st.session_state.rating
st.markdown(
    """
    <style>
        /* Estilos para el contenedor de valoración */
        .rating-container {
            display: flex;
            align-items: center;
            gap: 2rem;
            padding: 1.5rem;
            background-color: white;
            border-radius: 12px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        
        /* Estilos para los radio buttons */
        .stRadio > div {
            flex-direction: row !important;
            gap: 1rem !important;
        }
        
        .stRadio label {
            cursor: pointer !important;
            font-size: 1.5rem !important;
            padding: 0.5rem !important;
            border-radius: 8px !important;
            transition: all 0.2s ease !important;
        }
        
        .stRadio label:hover {
            background-color: rgba(45, 93, 160, 0.1) !important;
        }
        
        /* Ocultar el radio button original */
        .stRadio input[type="radio"] {
            position: absolute !important;
            opacity: 0 !important;
        }
        
        /* Estilo para la opción seleccionada */
        .stRadio input[type="radio"]:checked + label {
            color: #F59E0B !important;
            transform: scale(1.1) !important;
        }
    </style>
    """,
    unsafe_allow_html=True
)

def mostrar_ayuda():
    """Muestra la sección de ayuda."""
    st.header("Centro de ayuda")
    
    with st.expander("🔍 ¿Cómo funciona el asistente?"):
        st.write("""
        1. Selecciona tu área jurídica
        2. Elige el tipo de proceso
        3. Sube tus documentos
        4. Revisa las opciones sugeridas
        5. Descarga el documento generado
        """)
    
    with st.expander("📄 Tipos de documentos soportados"):
        st.write("- Archivos PDF (.pdf)\n- Documentos Word (.docx)")
    
    with st.expander("⚡ Consejos de uso"):
        st.write("""
        - Asegúrate de que los documentos estén bien escaneados
        - Revisa siempre el resultado antes de presentarlo
        - Guarda una copia de seguridad de tus documentos
        """)

def mostrar_contacto():
    """Muestra la información de contacto."""
    st.header("Contacto")
    
    # Franja descriptiva superior
    st.markdown(
        """
        <div class='descriptive-band'>
            Nuestro equipo de expertos está disponible para ayudarte con cualquier consulta legal
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(
            """
            <div class="contact-info">
                <h4>Soporte técnico</h4>
                <div class="contact-detail">
                    📧 melgarejobejarano@gmail.com
                </div>
                <div class="contact-detail">
                    📞 +51 961 277 144
                </div>
                <div style='font-size: 0.9em; color: var(--neutral-color); margin-top: 1rem;'>
                    Respuesta garantizada en menos de 24 horas hábiles
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    with col2:
        st.markdown(
            """
            <div class="contact-info">
                <h4>Horario de atención</h4>
                <div class="contact-detail">
                    🗓️ Lunes a Viernes
                </div>
                <div class="contact-detail">
                    ⏰ 9:00 AM - 6:00 PM
                </div>
                <div style='font-size: 0.9em; color: var(--neutral-color); margin-top: 1rem;'>
                    Hora de Perú (GMT-5)
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

###############################################################################
# Inicialización de estado                                                    
###############################################################################

# Inicializar estados de la aplicación
if 'page' not in st.session_state:
    st.session_state.page = "inicio"

if 'paso_actual' not in st.session_state:
    st.session_state.paso_actual = 1

if 'conversation_history' not in st.session_state:
    st.session_state.conversation_history = {
        'messages': [],
        'system_info': {
            'start_time': datetime.now().isoformat(),
            'openai_model': 'gpt-4',
        }
    }

if 'openai_thread' not in st.session_state:
    st.session_state.openai_thread = openai.beta.threads.create()

###############################################################################
# Funciones principales                                                       
###############################################################################

ASSISTANT_IDS = {
    "Derecho Civil": "asst_JEqVhFH9ertyrJTGFNq1zIZ0",
    "Derecho de Familia": "asst_k72lXItROiR9tgnDBiqmWf9j",
}

ROLES_PROCESALES = ["Demandante", "Demandado"]

# Palabras clave para determinar el formato
PALABRAS_FORMATO_ESCRITO = [
    "subsanar", "ampliar", "aclarar", "mero trámite",
    "téngase", "sírvase", "adjuntar", "presente"
]

PALABRAS_FORMATO_REGULAR = [
    "demanda", "denuncia", "apelación", "recurso",
    "impugnación", "nulidad", "casación", "queja"
]

openai.api_key = os.getenv("OPENAI_API_KEY")
if not openai.api_key:
    st.warning("⚠️  Defina la variable de entorno OPENAI_API_KEY para continuar.")

MAX_DOC_CHARS = 100_000  # límite para evitar desbordar tokens

def enviar_mensaje_y_esperar(mensaje: str, assistant_id: str) -> str:
    """Envía un mensaje al thread y espera la respuesta."""
    thread_id = st.session_state.openai_thread.id

    # Verificar y esperar si hay runs activos
    runs = openai.beta.threads.runs.list(thread_id=thread_id)
    for run in runs.data:
        if run.status not in {"completed", "failed", "cancelled", "expired"}:
            while True:
                time.sleep(1)
                run = openai.beta.threads.runs.retrieve(
                    thread_id=thread_id,
                    run_id=run.id
                )
                if run.status in {"completed", "failed", "cancelled", "expired"}:
                    break

    # Agregar el mensaje al thread existente
    openai.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content=mensaje
    )

    # Ejecutar el asistente
    run = openai.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_id
    )

    # Esperar la respuesta
    while run.status not in {"completed", "failed", "cancelled", "expired"}:
        time.sleep(1)
        run = openai.beta.threads.runs.retrieve(
            thread_id=thread_id,
            run_id=run.id
        )

    if run.status != "completed":
        return None

    # Obtener la última respuesta
    messages = openai.beta.threads.messages.list(
        thread_id=thread_id
    )
    return messages.data[0].content[0].text.value.strip()

def add_to_history(role: str, content: str, metadata: dict = None):
    """Agrega un mensaje al historial de conversación."""
    message = {
        'role': role,
        'content': content,
        'timestamp': datetime.now().isoformat(),
    }
    if metadata:
        message['metadata'] = metadata
    
    st.session_state.conversation_history['messages'].append(message)

def get_conversation_json() -> str:
    """Devuelve el historial de conversación en formato JSON."""
    return json.dumps(st.session_state.conversation_history, indent=2, ensure_ascii=False)

def _extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_text(uploaded_file) -> str:
    """Extrae texto de un PDF o Word subido."""
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".pdf":
        return _extract_text_from_pdf(uploaded_file)
    if suffix == ".docx":
        return _extract_text_from_docx(uploaded_file)
    raise ValueError("Formato no soportado: debe ser PDF o DOCX")

@st.cache_data(show_spinner=False, ttl=3600, max_entries=32)
def ai_analyze(document_text: str, assistant_id: str, area: str, rol: str) -> dict | None:
    """Envía el documento al asistente y obtiene la etapa procesal y soluciones."""
    # Reducir a tamaño manejable
    doc_chunk = document_text[:MAX_DOC_CHARS]

    # Registrar la consulta en el historial
    add_to_history('user', doc_chunk, {'area': area, 'assistant_id': assistant_id, 'rol': rol})

    mensaje = (
        f"Actúa como asistente jurídico en favor del {rol}. "
        "Analiza el siguiente documento legal y responde EXCLUSIVAMENTE "
                    "con un JSON que contenga las claves 'etapa_proceso' (string) "
        "y 'soluciones' (lista de EXACTAMENTE 3 strings numerados del 1-3, "
        f"que representen las mejores opciones legales para el {rol} "
        "en el siguiente paso procesal). No incluyas ningún texto adicional.\n\n"
                    f"Área de especialidad: {area}.\n\nDocumento:\n{doc_chunk}"
    )

    reply = enviar_mensaje_y_esperar(mensaje, assistant_id)
    if not reply:
        add_to_history('assistant', 'Error: No se pudo completar el análisis', {'status': 'failed'})
        return None

    # Registrar la respuesta en el historial
    add_to_history('assistant', reply, {'thread_id': st.session_state.openai_thread.id})

    try:
        data = json.loads(reply)
    except json.JSONDecodeError:
        # Intentar repararlo con un modelo de corrección rápida
        mensaje_correccion = "Corrige para que sea JSON válido sin comentarios.\n\n" + reply
        reply_corregido = enviar_mensaje_y_esperar(mensaje_correccion, assistant_id)
        data = json.loads(reply_corregido)

    return data

def determinar_formato(solucion: str) -> str:
    """Determina el formato necesario basado en el contenido de la solución."""
    solucion_lower = solucion.lower()
    
    for palabra in PALABRAS_FORMATO_ESCRITO:
        if palabra in solucion_lower:
            return "ESCRITO"
            
    for palabra in PALABRAS_FORMATO_REGULAR:
        if palabra in solucion_lower:
            return "REGULAR"
            
    # Por defecto, usar formato ESCRITO para trámites más simples
    return "ESCRITO"

def generar_prompt_redaccion(formato: str, solucion: str, area: str, rol: str, etapa: str) -> str:
    """Genera el prompt específico según el formato requerido."""
    
    instrucciones_comunes = """
    INSTRUCCIONES GENERALES:
    - Usa lenguaje claro y motivado (art. 122 CPC, art. 50 LOPJ)
    - Emplea conectores lógicos ("primero", "además", "por ende")
    - Mantén un tono empático pero firme
    - Vincula los hechos con derechos fundamentales afectados
    - Verifica plazos procesales
    - Cumple el Código de Ética del PJ y checklist CPC
    """
    
    if formato == "ESCRITO":
        template = f"""
        Redacta un ESCRITO JUDICIAL siguiendo EXACTAMENTE esta estructura:
        
        1. ENCABEZADO: órgano, expediente N°, materia (mayúsculas centradas)
        2. IDENTIFICACIÓN: datos completos de partes y apoderados (DNI/RUC, domicilio, casilla, poder)
        3. EXPOSICIÓN FÁCTICA: usa método PASTOR (Problem-Amplify-Story-Transformation-Offer-Request)
        4. FUNDAMENTOS DE DERECHO: normas (CPC, CC, Constitución) y precedentes con conectores lógicos
        5. PETITORIO: ítems numerados con verbos imperativos
        6. MEDIOS PROBATORIOS: lista ANEXO 1-n con descripción y folios
        7. FIRMA Y CIERRE: lugar, fecha, firma, CAPE, casilla
        
        Área: {area}
        Rol procesal: {rol}
        Etapa actual: {etapa}
        Solución elegida: {solucion}
        
        {instrucciones_comunes}
        """
    else:  # REGULAR
        template = f"""
        Redacta una PIEZA PROCESAL COMPLETA siguiendo EXACTAMENTE esta estructura:
        
        1. ENCABEZADO: órgano, tipo de acción, expediente
        2. PARTES Y APODERADOS: datos completos y legitimación
        3. COMPETENCIA: fundamento legal (territorial/cuantía)
        4. PETITORIO: pretensiones principales y subsidiarias valoradas
        5. FUNDAMENTOS DE HECHO: cronología con cierre jurídico
        6. FUNDAMENTOS DE DERECHO: artículos y precedentes (incluir arts. 2-3-139 Const.)
        7. REQUISITOS DE ADMISIBILIDAD: checklist CPC 130-135
        8. MEDIOS PROBATORIOS: descripción + finalidad probatoria
        9. ANEXOS: documentos numerados y foliados
        10. PLAZO/AGRAVIO: cómputo detallado si aplica
        11. PETICIÓN Y COSTAS: art. 56 CPC
        12. FIRMA Y CIERRE: datos completos con CAPE
        
        Área: {area}
        Rol procesal: {rol}
        Etapa actual: {etapa}
        Solución elegida: {solucion}
        
        {instrucciones_comunes}
        """
    
    return template

def ai_draft(solution: str, stage: str, assistant_id: str, area: str, rol: str, original_text: str) -> str:
    """Solicita al asistente la redacción del escrito judicial."""
    
    # Determinar el formato necesario
    formato = determinar_formato(solution)
    
    # Generar el prompt específico
    prompt = generar_prompt_redaccion(formato, solution, area, rol, stage)
    
    # Registrar la solicitud de redacción
    add_to_history('user', prompt, {
        'area': area,
        'stage': stage,
        'solution': solution,
        'formato': formato,
        'rol': rol,
        'assistant_id': assistant_id
    })

    response = enviar_mensaje_y_esperar(prompt, assistant_id)
    if not response:
        return "Error: No se pudo generar el documento."

    # Registrar la respuesta
    add_to_history('assistant', response, {
        'thread_id': st.session_state.openai_thread.id,
        'formato': formato
    })

    return response

def generar_historial():
    """Genera un historial completo de documentos."""
    if not st.session_state.conversation_history['messages']:
        st.warning("No hay historial de documentos disponible.")
        return None
        
    historial = {
        'fecha_generacion': datetime.now().isoformat(),
        'documentos': []
    }
    
    for msg in st.session_state.conversation_history['messages']:
        if msg['role'] == 'assistant' and 'formato' in msg.get('metadata', {}):
            historial['documentos'].append({
                'timestamp': msg['timestamp'],
                'contenido': msg['content'],
                'formato': msg['metadata']['formato']
            })
    
    return historial

###############################################################################
# Interfaz principal                                                           
###############################################################################

# Mostrar navegación
mostrar_navegacion()

# Contenido según la página actual
if st.session_state.page == "inicio":
    st.title("Experto Jurídico 7.0")

    # Logo y descripción en diseño responsivo
    col1, col2 = st.columns([3, 2], gap="large")
    
    with col1:
        st.write("""
        Bienvenido al asistente jurídico inteligente que te ayuda a generar
        documentos legales profesionales de manera rápida y precisa.
        """)
        
        st.markdown(
            """
            <div class='descriptive-band'>
                Potenciado por inteligencia artificial para brindarte la mejor asesoría legal
            </div>
            """,
            unsafe_allow_html=True
        )
        
        st.markdown(
            """
            <div class='info-box'>
                💡 Nuestro asistente utiliza tecnología de punta para:
                <ul>
                    <li>Analizar documentos legales</li>
                    <li>Sugerir estrategias procesales</li>
                    <li>Generar escritos profesionales</li>
                    <li>Optimizar tu tiempo y recursos</li>
                </ul>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        if st.button("Comenzar ▶️", type="primary"):
            st.session_state.page = "generar"
            
    with col2:
        image_path = generar_imagen_corporativa()
        if image_path:
            st.markdown(
                f"""
                <div class="corporate-image-container">
                    <img src="data:image/png;base64,{base64.b64encode(open(image_path, 'rb').read()).decode()}" 
                         alt="Asistente Jurídico">
                </div>
                """,
                unsafe_allow_html=True
            )

elif st.session_state.page == "generar":
    st.title("Generación de documentos")
    
    # Inicializar paso_actual si no existe
    if 'paso_actual' not in st.session_state:
        st.session_state.paso_actual = 1
    
    # Progreso del asistente
    progress = (st.session_state.paso_actual - 1) * 25
    st.progress(progress)
    
    # Paso 1: Especialidad jurídica
    if st.session_state.paso_actual == 1:
        st.header("Paso 1: Especialidad jurídica")
        area = st.radio(
            "Seleccione su área de especialidad:",
            ["⚖️ Derecho Civil", "👨‍👩‍👧‍👦 Derecho de Familia"],
            format_func=lambda x: x.split(" ", 1)[1],
            horizontal=True
        )
        if st.button("Continuar ▶️", type="primary") and area:
            st.session_state.area = area.split(" ", 1)[1]
            st.session_state.paso_actual = 2
            st.rerun()
    
    # Paso 2: Rol procesal
    elif st.session_state.paso_actual == 2:
        st.header("Paso 2: Rol procesal")
        rol = st.radio(
            "Seleccione su rol en el proceso:",
            ["⚔️ Demandante", "🛡️ Demandado"],
            format_func=lambda x: x.split(" ", 1)[1],
            horizontal=True
        )
        if st.button("Continuar ▶️", type="primary") and rol:
            st.session_state.rol = rol.split(" ", 1)[1]
            st.session_state.paso_actual = 3
            st.rerun()
    
    # Paso 3: Documentos
    elif st.session_state.paso_actual == 3:
        st.header("Paso 3: Documentos")
        uploaded_file = st.file_uploader(
            "Subir documento",
            type=["pdf", "docx"],
            help="El documento debe estar en formato PDF o Word (.docx)",
            label_visibility="collapsed"
        )
        
        if uploaded_file:
            # Mostrar información del archivo
            file_details = {
                "Nombre": uploaded_file.name,
                "Tipo": uploaded_file.type,
                "Tamaño": f"{uploaded_file.size / 1024:.1f} KB"
            }
            
            st.info("Documento cargado correctamente")
            for key, value in file_details.items():
                st.text(f"{key}: {value}")
            
            if st.button("Analizar documento ▶️", type="primary"):
                try:
                    with st.spinner("Analizando documento..."):
                        text = extract_text(uploaded_file)
                        assistant_id = ASSISTANT_IDS[st.session_state.area]
                        analysis = ai_analyze(text, assistant_id, st.session_state.area, st.session_state.rol)
                        
                        if analysis:
                            st.session_state.analysis = analysis
                            st.session_state.document_text = text
                            st.session_state.paso_actual = 4
                            st.toast("¡Análisis completado! 🎉")
                            st.rerun()
                        else:
                            st.error("No se pudo analizar el documento.")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
        else:
            st.info("👆 Sube un documento para continuar")

# Paso 4: Opciones
elif st.session_state.paso_actual == 4:
    st.header("Paso 4: Opciones sugeridas")
    analysis = st.session_state.analysis
    st.info(f"**Etapa procesal:** {analysis['etapa_proceso']}")
    choice = st.radio(
        "Seleccione la acción a realizar:",
        analysis['soluciones'],
        format_func=lambda x: x.strip("123. "),
        horizontal=True
    )
    if choice:
        formato = determinar_formato(choice)
        st.info(f"**Formato sugerido:** {formato}")
    if st.button("Generar documento ▶️", type="primary") and choice:
        with st.spinner("Generando documento..."):
            draft_text = ai_draft(
                choice,
                analysis['etapa_proceso'],
                ASSISTANT_IDS[st.session_state.area],
                st.session_state.area,
                st.session_state.rol,
                st.session_state.document_text
            )
            st.session_state.draft_text = draft_text
            st.session_state.paso_actual = 5
            st.toast("¡Documento generado! 📄")
            st.rerun()

# Paso 5: Resultado
elif st.session_state.paso_actual == 5:
    st.header("Paso 5: Documento generado")
    with st.expander("Ver documento", expanded=True):
        st.text_area(
            "Contenido del documento",
            st.session_state.draft_text,
            height=500,
            help="Puede copiar el texto o descargarlo en formato Word"
        )
    col1, col2 = st.columns(2)
    with col1:
        # Generar documento Word
        doc = Document()
        doc.add_paragraph(st.session_state.draft_text)
        # Guardar temporalmente
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            # Leer el archivo para la descarga
            with open(tmp.name, "rb") as docx_file:
                docx_bytes = docx_file.read()
                st.download_button(
                    label="📥 Descargar DOCX",
                    data=docx_bytes,
                    file_name="documento_legal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            # Limpiar archivo temporal
            os.unlink(tmp.name)
    
    with col2:
        # Botón para descargar el historial
        conversation_json = get_conversation_json()
        st.download_button(
            label="📥 Descargar historial",
            data=conversation_json,
            file_name=f"historial_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json"
        )
    
    if st.button("Nuevo documento 🔄", type="primary"):
        st.session_state.paso_actual = 1
        st.rerun()

elif st.session_state.page == "historial":
    st.title("Historial de documentos")
    
    st.markdown(
        """
        <div class='descriptive-band'>
            Accede al historial completo de documentos generados
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Botón para generar historial
    st.markdown(
        """
        <button class="history-button" onclick="document.dispatchEvent(new CustomEvent('generate-history'))">
            📋 Generar historial
        </button>
        """,
        unsafe_allow_html=True
    )
    
    if st.session_state.get('show_history'):
        historial = generar_historial()
        if historial:
            st.json(historial)
            
            # Botón de descarga
            st.download_button(
                label="📥 Descargar historial",
                data=json.dumps(historial, indent=2, ensure_ascii=False),
                file_name=f"historial_documentos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )

elif st.session_state.page == "ayuda":
    mostrar_ayuda()

elif st.session_state.page == "contacto":
    mostrar_contacto()

elif st.session_state.page == "feedback":
    mostrar_feedback() 

# Actualizar el manejo del estado en Streamlit
if 'rating_value' not in st.session_state:
    st.session_state.rating_value = 0

# Agregar un callback para manejar el cambio de rating
def handle_rating_change():
    st.session_state.rating_changed = True

# Componente oculto para manejar la actualización del estado
st.markdown(
    f"""
    <div style="display: none;">
        <input type="number" 
               value="{st.session_state.rating_value}" 
               key="rating_input" 
               on_change="handle_rating_change"
        />
    </div>
    """,
    unsafe_allow_html=True
) 